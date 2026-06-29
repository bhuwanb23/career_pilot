import json
from pathlib import Path
from unittest.mock import AsyncMock, MagicMock, patch

import pytest

SAMPLE_PDF = Path(__file__).resolve().parent.parent / "samples" / "sample_resume_john.pdf"

MOCK_ENHANCED_RESPONSE = json.dumps({
    "personal": {
        "name": "John Doe",
        "email": "john@example.com",
        "phone": "555-0123",
        "location": "San Francisco, CA",
        "linkedin": "linkedin.com/in/johndoe",
        "github": "github.com/johndoe",
    },
    "summary": "Experienced software engineer with 5+ years in full-stack development.",
    "skills": ["Python", "JavaScript", "React", "Node.js", "SQL", "FastAPI", "Git", "Docker"],
    "projects": [
        {"name": "CareerPilot", "description": "AI career assistant", "tech": ["Python", "FastAPI", "React"]}
    ],
    "education": [
        {"school": "UC Berkeley", "degree": "BS", "field": "Computer Science", "year": "2020", "gpa": "3.8"}
    ],
    "experience": [
        {"company": "TechCorp", "role": "Software Engineer", "start_date": "01/2022", "end_date": "12/2024", "bullets": ["Built REST APIs", "Led microservices migration"]}
    ],
    "certifications": [
        {"name": "AWS Solutions Architect", "issuer": "Amazon", "year": "2023"}
    ],
    "languages": [
        {"language": "English", "proficiency": "Native"},
        {"language": "Spanish", "proficiency": "Intermediate"}
    ],
})

MOCK_MINIMAL_RESPONSE = json.dumps({
    "personal": {"name": "Jane Smith", "email": "jane@test.com", "phone": "", "location": "", "linkedin": "", "github": ""},
    "summary": "Data scientist",
    "skills": ["Python", "ML"],
    "projects": [],
    "education": [],
    "experience": [],
    "certifications": [],
    "languages": [],
})


@pytest.fixture(autouse=True)
def reset_llm():
    import services.llm_client as lc
    lc._provider = None
    yield
    lc._provider = None


@pytest.fixture
def mock_llm():
    mock_provider = MagicMock()
    mock_provider.generate = AsyncMock()
    with patch("services.llm_client.get_llm_provider", return_value=mock_provider):
        yield mock_provider


class TestParsedResumeValidation:
    def test_validates_complete_data(self):
        from services.resume_parser import ParsedResume
        data = json.loads(MOCK_ENHANCED_RESPONSE)
        parsed = ParsedResume(**data)
        assert parsed.personal["name"] == "John Doe"
        assert parsed.personal["email"] == "john@example.com"
        assert len(parsed.skills) == 8
        assert len(parsed.experience) == 1
        assert len(parsed.certifications) == 1
        assert len(parsed.languages) == 2

    def test_validates_minimal_data(self):
        from services.resume_parser import ParsedResume
        data = json.loads(MOCK_MINIMAL_RESPONSE)
        parsed = ParsedResume(**data)
        assert parsed.personal["name"] == "Jane Smith"
        assert len(parsed.skills) == 2
        assert len(parsed.certifications) == 0

    def test_skills_string_normalization(self):
        from services.resume_parser import ParsedResume
        parsed = ParsedResume(skills="Python, JavaScript, React")
        assert parsed.skills == ["Python", "JavaScript", "React"]

    def test_experience_string_normalization(self):
        from services.resume_parser import ParsedResume
        parsed = ParsedResume(experience=["Software Engineer"])
        assert parsed.experience[0]["role"] == "Software Engineer"

    def test_education_string_normalization(self):
        from services.resume_parser import ParsedResume
        parsed = ParsedResume(education=["UC Berkeley"])
        assert parsed.education[0]["school"] == "UC Berkeley"

    def test_certifications_string_normalization(self):
        from services.resume_parser import ParsedResume
        parsed = ParsedResume(certifications=["AWS Certified"])
        assert parsed.certifications[0]["name"] == "AWS Certified"

    def test_languages_string_normalization(self):
        from services.resume_parser import ParsedResume
        parsed = ParsedResume(languages=["English"])
        assert parsed.languages[0]["language"] == "English"

    def test_empty_personal(self):
        from services.resume_parser import ParsedResume
        parsed = ParsedResume(personal={})
        assert parsed.personal == {}

    def test_empty_lists_default(self):
        from services.resume_parser import ParsedResume
        parsed = ParsedResume()
        assert parsed.skills == []
        assert parsed.projects == []
        assert parsed.education == []
        assert parsed.experience == []
        assert parsed.certifications == []
        assert parsed.languages == []


class TestProfileService:
    def test_create_profile_with_personal_info(self, db_session):
        from services.profile_service import create_or_update_profile, get_profile
        data = {
            "raw_resume": "resume text",
            "personal": {"name": "John", "email": "john@test.com", "phone": "555-0123"},
            "summary": "Developer",
            "skills": ["Python"],
            "certifications": [{"name": "AWS", "issuer": "Amazon", "year": "2023"}],
            "languages": [{"language": "English", "proficiency": "Native"}],
        }
        profile = create_or_update_profile(db_session, data)
        assert profile.personal_name == "John"
        assert profile.personal_email == "john@test.com"
        assert profile.personal_phone == "555-0123"
        assert profile.get_certifications() == [{"name": "AWS", "issuer": "Amazon", "year": "2023"}]
        assert profile.get_languages() == [{"language": "English", "proficiency": "Native"}]

    def test_profile_to_dict_includes_new_fields(self, db_session):
        from services.profile_service import create_or_update_profile, profile_to_dict
        data = {
            "personal": {"name": "Jane", "email": "jane@test.com"},
            "summary": "Data scientist",
            "skills": ["Python", "ML"],
            "certifications": [{"name": "GCP", "issuer": "Google", "year": "2024"}],
            "languages": [{"language": "Spanish", "proficiency": "Fluent"}],
        }
        profile = create_or_update_profile(db_session, data)
        d = profile_to_dict(profile)
        assert d["personal"]["name"] == "Jane"
        assert d["personal"]["email"] == "jane@test.com"
        assert len(d["certifications"]) == 1
        assert len(d["languages"]) == 1

    def test_update_existing_profile(self, db_session):
        from services.profile_service import create_or_update_profile, get_profile
        create_or_update_profile(db_session, {"summary": "First", "skills": ["A"]})
        create_or_update_profile(db_session, {"summary": "Updated", "skills": ["A", "B"]})
        profile = get_profile(db_session)
        assert profile.summary == "Updated"
        assert profile.get_skills() == ["A", "B"]


class TestParseResume:
    @pytest.mark.asyncio
    async def test_parse_resume_enhanced_output(self, mock_llm):
        mock_llm.generate.return_value = MOCK_ENHANCED_RESPONSE
        with open(SAMPLE_PDF, "rb") as f:
            content = f.read()
        from services.resume_parser import parse_resume
        result = await parse_resume(content, "resume.pdf")
        assert result["raw_resume"]
        assert result["personal"]["name"] == "John Doe"
        assert result["personal"]["email"] == "john@example.com"
        assert len(result["skills"]) == 8
        assert len(result["certifications"]) == 1
        assert len(result["languages"]) == 2

    @pytest.mark.asyncio
    async def test_parse_resume_fallback_on_bad_json(self, mock_llm):
        mock_llm.generate.return_value = "not valid json"
        with open(SAMPLE_PDF, "rb") as f:
            content = f.read()
        from services.resume_parser import parse_resume
        result = await parse_resume(content, "resume.pdf")
        assert "raw_resume" in result
        assert result["skills"] == []

    @pytest.mark.asyncio
    async def test_parse_resume_empty_text_raises(self, mock_llm):
        from docx import Document
        import io
        doc = Document()
        buf = io.BytesIO()
        doc.save(buf)
        from services.resume_parser import parse_resume
        with pytest.raises(ValueError, match="No text"):
            await parse_resume(buf.getvalue(), "empty.docx")


class TestDocumentExtractor:
    @pytest.mark.asyncio
    async def test_extract_pdf(self):
        from services.document_extractor import extract_document
        with open(SAMPLE_PDF, "rb") as f:
            content = f.read()
        result = await extract_document(content, "resume.pdf")
        assert "text" in result
        assert len(result["text"]) > 0
        assert result["engine"] == "pymupdf"

    @pytest.mark.asyncio
    async def test_extract_docx(self):
        from services.document_extractor import extract_document
        from docx import Document
        import io
        doc = Document()
        doc.add_paragraph("John Doe")
        doc.add_paragraph("Python Developer")
        buf = io.BytesIO()
        doc.save(buf)
        result = await extract_document(buf.getvalue(), "resume.docx")
        assert "John Doe" in result["text"]
        assert result["engine"] == "python-docx"

    @pytest.mark.asyncio
    async def test_extract_unsupported_returns_empty(self):
        from services.document_extractor import extract_document
        result = await extract_document(b"hello", "file.txt")
        assert result["text"] == ""


class TestPersonalInfoExtraction:
    @pytest.mark.asyncio
    async def test_personal_info_in_profile(self, mock_llm):
        mock_llm.generate.return_value = MOCK_ENHANCED_RESPONSE
        with open(SAMPLE_PDF, "rb") as f:
            content = f.read()
        from services.resume_parser import parse_resume
        from services.profile_service import create_or_update_profile
        from database import SessionLocal
        result = await parse_resume(content, "resume.pdf")
        db = SessionLocal()
        profile = create_or_update_profile(db, result)
        assert profile.personal_name == "John Doe"
        assert profile.personal_email == "john@example.com"
        assert profile.personal_phone == "555-0123"
        assert profile.personal_location == "San Francisco, CA"
        assert profile.personal_linkedin == "linkedin.com/in/johndoe"
        assert profile.personal_github == "github.com/johndoe"
        db.close()

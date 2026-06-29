import io
from pathlib import Path
from unittest.mock import AsyncMock, MagicMock, patch

import pytest

SAMPLE_PDF = Path(__file__).resolve().parent.parent / "samples" / "sample_resume_john.pdf"

MOCK_RESUME_RESPONSE = '{"summary": "Experienced developer", "skills": ["Python", "React"], "projects": [], "education": [{"school": "UC Berkeley", "degree": "BS CS", "year": "2020"}], "experience": [{"company": "TechCorp", "role": "SE", "dates": "2022-2024", "bullets": ["Built APIs"]}]}'


@pytest.fixture(autouse=True)
def reset_llm():
    import services.llm_client as lc
    lc._provider = None
    yield
    lc._provider = None


@pytest.fixture
def mock_llm():
    mock_provider = MagicMock()
    mock_provider.generate = AsyncMock()
    mock_provider.health_check = AsyncMock(return_value=True)
    with patch("services.llm_client.get_llm_provider", return_value=mock_provider):
        yield mock_provider


def _make_docx(text="John Doe\nPython Developer\nSkills: Python, React"):
    from docx import Document
    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestPDFUpload:
    def test_upload_pdf_success(self, client, mock_llm):
        mock_llm.generate.return_value = MOCK_RESUME_RESPONSE
        with open(SAMPLE_PDF, "rb") as f:
            r = client.post("/api/resume/upload", files={"file": ("resume.pdf", f, "application/pdf")})
        assert r.status_code == 200
        data = r.json()
        assert data["upload_id"] is not None
        assert data["filename"] == "resume.pdf"
        assert data["file_type"] == "pdf"
        assert data["status"] == "parsed"
        assert data["profile_id"] is not None

    def test_upload_creates_resume_upload_record(self, client, mock_llm):
        mock_llm.generate.return_value = MOCK_RESUME_RESPONSE
        with open(SAMPLE_PDF, "rb") as f:
            r = client.post("/api/resume/upload", files={"file": ("resume.pdf", f, "application/pdf")})
        upload_id = r.json()["upload_id"]
        r = client.get(f"/api/resume/uploads/{upload_id}")
        assert r.status_code == 200
        assert r.json()["status"] == "parsed"

    def test_upload_stores_file_on_disk(self, client, mock_llm):
        mock_llm.generate.return_value = MOCK_RESUME_RESPONSE
        with open(SAMPLE_PDF, "rb") as f:
            r = client.post("/api/resume/upload", files={"file": ("resume.pdf", f, "application/pdf")})
        upload_id = r.json()["upload_id"]
        from database import SessionLocal
        from models import ResumeUpload
        db = SessionLocal()
        upload = db.query(ResumeUpload).get(upload_id)
        assert upload.storage_path is not None
        from services.file_storage import get_stored_file
        content = get_stored_file(upload.storage_path)
        assert content is not None
        assert len(content) > 0
        db.close()


class TestDOCXUpload:
    def test_upload_docx_success(self, client, mock_llm):
        mock_llm.generate.return_value = MOCK_RESUME_RESPONSE
        docx_content = _make_docx()
        r = client.post("/api/resume/upload", files={"file": ("resume.docx", docx_content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")})
        assert r.status_code == 200
        data = r.json()
        assert data["file_type"] == "docx"
        assert data["status"] == "parsed"

    def test_upload_docx_creates_record(self, client, mock_llm):
        mock_llm.generate.return_value = MOCK_RESUME_RESPONSE
        docx_content = _make_docx("Jane Smith\nML Engineer")
        r = client.post("/api/resume/upload", files={"file": ("jane.docx", docx_content, "application/octet-stream")})
        upload_id = r.json()["upload_id"]
        r = client.get(f"/api/resume/uploads/{upload_id}")
        assert r.status_code == 200
        assert r.json()["filename"] == "jane.docx"


class TestValidation:
    def test_reject_unsupported_file_type(self, client):
        r = client.post("/api/resume/upload", files={"file": ("resume.txt", b"hello", "text/plain")})
        assert r.status_code == 400
        assert "Unsupported" in r.json()["detail"]

    def test_reject_exe_renamed_to_pdf(self, client):
        r = client.post("/api/resume/upload", files={"file": ("fake.pdf", b"MZ\x90\x00", "application/pdf")})
        assert r.status_code in (400, 500)

    def test_reject_empty_file(self, client):
        r = client.post("/api/resume/upload", files={"file": ("empty.pdf", b"", "application/pdf")})
        assert r.status_code == 400
        assert "empty" in r.json()["detail"].lower()

    def test_reject_oversized_file(self, client):
        big_content = b"x" * (6 * 1024 * 1024)
        r = client.post("/api/resume/upload", files={"file": ("big.pdf", big_content, "application/pdf")})
        assert r.status_code == 413


class TestUploadHistory:
    def test_list_uploads_empty(self, client):
        r = client.get("/api/resume/uploads")
        assert r.status_code == 200
        assert r.json() == []

    def test_list_uploads_after_upload(self, client, mock_llm):
        mock_llm.generate.return_value = MOCK_RESUME_RESPONSE
        with open(SAMPLE_PDF, "rb") as f:
            client.post("/api/resume/upload", files={"file": ("resume.pdf", f, "application/pdf")})
        r = client.get("/api/resume/uploads")
        assert r.status_code == 200
        assert len(r.json()) == 1
        assert r.json()[0]["filename"] == "resume.pdf"

    def test_get_upload_detail(self, client, mock_llm):
        mock_llm.generate.return_value = MOCK_RESUME_RESPONSE
        with open(SAMPLE_PDF, "rb") as f:
            r = client.post("/api/resume/upload", files={"file": ("resume.pdf", f, "application/pdf")})
        upload_id = r.json()["upload_id"]
        r = client.get(f"/api/resume/uploads/{upload_id}")
        assert r.status_code == 200
        data = r.json()
        assert "raw_text_preview" in data
        assert len(data["raw_text_preview"]) > 0

    def test_get_upload_404(self, client):
        r = client.get("/api/resume/uploads/9999")
        assert r.status_code == 404


class TestDownload:
    def test_download_upload(self, client, mock_llm):
        mock_llm.generate.return_value = MOCK_RESUME_RESPONSE
        with open(SAMPLE_PDF, "rb") as f:
            r = client.post("/api/resume/upload", files={"file": ("resume.pdf", f, "application/pdf")})
        upload_id = r.json()["upload_id"]
        r = client.get(f"/api/resume/uploads/{upload_id}/download")
        assert r.status_code == 200
        assert r.headers["content-type"] == "application/pdf"
        assert len(r.content) > 0

    def test_download_404(self, client):
        r = client.get("/api/resume/uploads/9999/download")
        assert r.status_code == 404


class TestDelete:
    def test_delete_upload(self, client, mock_llm):
        mock_llm.generate.return_value = MOCK_RESUME_RESPONSE
        with open(SAMPLE_PDF, "rb") as f:
            r = client.post("/api/resume/upload", files={"file": ("resume.pdf", f, "application/pdf")})
        upload_id = r.json()["upload_id"]
        r = client.delete(f"/api/resume/uploads/{upload_id}")
        assert r.status_code == 200
        r = client.get("/api/resume/uploads")
        assert len(r.json()) == 0

    def test_delete_404(self, client):
        r = client.delete("/api/resume/uploads/9999")
        assert r.status_code == 404


class TestFileStorage:
    def test_store_and_retrieve(self):
        from services.file_storage import store_file, get_stored_file, delete_stored_file
        content = b"test content"
        path = store_file("test.pdf", content)
        assert path.startswith("uploads/")
        retrieved = get_stored_file(path)
        assert retrieved == content
        assert delete_stored_file(path)
        assert get_stored_file(path) is None

    def test_compute_hash(self):
        from services.file_storage import compute_file_hash
        h = compute_file_hash(b"hello")
        assert len(h) == 64
        assert h == compute_file_hash(b"hello")


class TestDOCXExtraction:
    def test_extract_docx(self):
        from services.document_extractor import _extract_with_docx
        content = _make_docx("John Doe\nPython Developer\nSkills: Python, React")
        result = _extract_with_docx(content)
        assert "John Doe" in result["text"]
        assert "Python Developer" in result["text"]
        assert result["engine"] == "python-docx"

    def test_extract_docx_empty(self):
        from services.document_extractor import _extract_with_docx
        from docx import Document
        doc = Document()
        buf = io.BytesIO()
        doc.save(buf)
        result = _extract_with_docx(buf.getvalue())
        assert result["text"] == ""


class TestPipelineIntegration:
    def test_upload_advances_pipeline(self, client, mock_llm):
        from models import Application
        from database import SessionLocal
        db = SessionLocal()
        app = Application(company="Test", role="Dev", job_description="jd")
        db.add(app)
        db.commit()
        app_id = app.id
        db.close()

        mock_llm.generate.return_value = MOCK_RESUME_RESPONSE
        with open(SAMPLE_PDF, "rb") as f:
            client.post("/api/resume/upload", files={"file": ("resume.pdf", f, "application/pdf")})

        db = SessionLocal()
        from models import ApplicationPipeline
        pipeline = db.query(ApplicationPipeline).filter_by(application_id=app_id).first()
        assert pipeline is not None
        assert pipeline.current_stage == "resume_uploaded"
        db.close()
